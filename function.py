# functions.py

import os
import json
import time
import requests
import re
from pathlib import Path
from PySide6.QtWidgets import QFileDialog
from youtube_transcript_api import (
    YouTubeTranscriptApi,
    NoTranscriptFound,
    TranscriptsDisabled,
    VideoUnavailable,
)
from json import JSONDecodeError

# -------------------------
# Helper: read_file_with_fallback
# -------------------------
def read_file_with_fallback(filepath):
    encodings = ["utf-8", "latin1", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError(f"Could not decode file {filepath} with available encodings.")


# -------------------------
# Ollama API Helper Function
# -------------------------
def generate_response(prompt, model, host="http://localhost:11434", cancel_event=None):
    url = f"{host}/api/generate"
    payload = {"model": model, "prompt": prompt, "stream": False}
    headers = {"Content-Type": "application/json"}
    try:
        if cancel_event and cancel_event.is_set():
            return "[Generation cancelled]", None
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
        if cancel_event and cancel_event.is_set():
            return "[Generation cancelled]", None
        response.raise_for_status()
        try:
            json_response = response.json()
        except JSONDecodeError:
            # Retry once if empty response
            time.sleep(1)
            response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
            response.raise_for_status()
            json_response = response.json()
        generated_text = json_response.get("response", "").strip()
        return generated_text, json_response
    except Exception as e:
        msg = str(e)
        if "no element found" in msg:
            return "[Unable to fetch transcript or response is empty]", None
        return f"[Error processing chunk: {e}]", None


# -------------------------
# Configuration
# -------------------------
class Config:
    def __init__(self):
        self.base_dir = Path(__file__).parent
        self.config_file = self.base_dir / "config.json"
        self.output_dir = self.base_dir / "outputs"
        self.temp_dir = self.base_dir / "temp"
        self.history_file = self.base_dir / "history.json"
        self._init_directories()
        self.settings = self._load_config()

    def _init_directories(self):
        self.output_dir.mkdir(exist_ok=True)
        self.temp_dir.mkdir(exist_ok=True)
        (self.temp_dir / "yt_trans").mkdir(exist_ok=True)
        (self.temp_dir / "yt_chunks").mkdir(exist_ok=True)
        (self.temp_dir / "yt_pro").mkdir(exist_ok=True)
        if not self.history_file.exists():
            with open(self.history_file, "w") as f:
                json.dump([], f)

    def _load_config(self):
        defaults = {
            "chunk_size": 300,
            "chunk_overlap": 50,
            "ollama_model": "deepseek-r1",
            "processing_prompt": "Check and reformat the text for grammar, clarity, and proper structure.",
            "output_format": "docx",
            "skip_manual_name": False,
            "last_video_id": "",
            "inline_output_name": "",
            "include_docx_title": True,
            "title_font_size": 16,
            "custom_title": "",
            "retry_count": 3,
            "typewriter_speed": 2,
        }
        try:
            if self.config_file.exists():
                with open(self.config_file, "r") as f:
                    data = json.load(f)
                    return {**defaults, **data}
        except JSONDecodeError:
            pass
        return defaults

    def save_config(self):
        with open(self.config_file, "w") as f:
            json.dump(self.settings, f, indent=2)

    def add_to_history(self, video_id, url, title=""):
        try:
            try:
                with open(self.history_file, "r", encoding="utf-8") as f:
                    history = json.load(f)
            except (FileNotFoundError, JSONDecodeError):
                history = []

            entry = {
                "id": video_id,
                "url": url,
                "title": title,
                "date": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            history.insert(0, entry)
            history = history[:50]
            with open(self.history_file, "w", encoding="utf-8") as f:
                json.dump(history, f, indent=2)
        except Exception as e:
            print(f"Error saving history: {e}")

    def load_history(self):
        try:
            with open(self.history_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except (FileNotFoundError, JSONDecodeError):
            return []

    def clean_temp(self):
        for subdir in ["yt_trans", "yt_chunks", "yt_pro"]:
            dir_path = self.temp_dir / subdir
            for item in dir_path.glob("*"):
                if item.is_file():
                    item.unlink()


# -------------------------
# Transcript Handling
# -------------------------
class TranscriptHandler:
    def __init__(self, config: Config):
        self.config = config
        self.config.clean_temp()

    def extract_and_save_transcript(self, video_url):
        retry_count = int(self.config.settings.get("retry_count", 3))
        retry_delay = 1

        for attempt in range(retry_count + 1):
            try:
                if "youtu.be" in video_url:
                    video_id = video_url.split("/")[-1]
                else:
                    match = re.search(r"v=([a-zA-Z0-9_-]+)", video_url)
                    if match:
                        video_id = match.group(1)
                    else:
                        raise ValueError("Invalid YouTube URL format")

                transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
                transcript_text = "\n".join([entry["text"] for entry in transcript_list])
                trans_dir = self.config.temp_dir / "yt_trans"
                transcript_file = trans_dir / f"{video_id}_transcript.txt"
                transcript_file.write_text(transcript_text, encoding="utf-8")
                self.config.settings["last_video_id"] = video_id
                self.config.save_config()

                video_title = self.get_youtube_title(video_id)
                self.config.add_to_history(video_id, video_url, video_title)

                return transcript_file, video_id, video_title
            except (NoTranscriptFound, TranscriptsDisabled, VideoUnavailable):
                if attempt < retry_count:
                    time.sleep(retry_delay)
                    retry_delay *= 2
                    continue
                raise RuntimeError("Transcript unavailable for this video.")
            except Exception as e:
                if attempt < retry_count:
                    time.sleep(retry_delay)
                    retry_delay *= 2
                    continue
                raise RuntimeError(f"Error extracting transcript: {e}")

    def get_youtube_title(self, video_id):
        url = f"https://www.youtube.com/watch?v={video_id}"
        try:
            response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=10)
            if response.status_code == 200:
                pattern = r'<title>(.*?)</title>'
                match = re.search(pattern, response.text)
                if match:
                    title = match.group(1).replace(' - YouTube', '').strip()
                    return title
        except Exception:
            pass
        return f"Video-{video_id}"

    def split_transcript(self, transcript_file):
        try:
            chunk_size = int(self.config.settings.get("chunk_size", 300))
            chunk_overlap = int(self.config.settings.get("chunk_overlap", 50))
            content = transcript_file.read_text(encoding="utf-8")
            words = content.split()
            total_words = len(words)
            chunks_dir = self.config.temp_dir / "yt_chunks"
            chunk_files = []
            start = 0
            chunk_id = 1

            while start < total_words:
                end = min(start + chunk_size, total_words)
                chunk_text = " ".join(words[start:end])
                chunk_file = chunks_dir / f"chunk_{chunk_id}.txt"
                chunk_file.write_text(chunk_text, encoding="utf-8")
                chunk_files.append(chunk_file)
                start += chunk_size - chunk_overlap
                chunk_id += 1

            return chunk_files
        except Exception as e:
            raise RuntimeError(f"Error splitting transcript: {e}")

    def process_single_chunk(self, chunk_file, cancel_event=None):
        try:
            chunk_content = chunk_file.read_text(encoding="utf-8")
            processing_prompt = self.config.settings.get(
                "processing_prompt",
                "Check and reformat the text for grammar, clarity, and proper structure.",
            )
            combined_prompt = (
                f"Processing Instruction:\n{processing_prompt}\n\n"
                f"Apply the above instruction to the following text:\n{chunk_content}"
            )
            generated_text, _ = generate_response(
                combined_prompt,
                self.config.settings.get("ollama_model", "deepseek-r1"),
                cancel_event=cancel_event,
            )
            output_dir = self.config.temp_dir / "yt_pro"
            output_dir.mkdir(exist_ok=True)
            output_file = output_dir / chunk_file.name
            output_file.write_text(generated_text, encoding="utf-8")
            return generated_text
        except Exception as e:
            return f"[Error processing chunk: {e}]"

    def combine_chunks_to_output(self, video_id, video_title="", status_callback=None):
        processed_dir = self.config.temp_dir / "yt_pro"
        processed_files = sorted(processed_dir.glob("*.txt"))
        if not processed_files:
            if status_callback:
                status_callback("Error: No processed chunks to combine.", "#ff7373")
            return

        if self.config.settings.get("skip_manual_name", False):
            default_name = video_id
        else:
            default_name = self.config.settings.get("inline_output_name", "").strip() or video_id

        output_format = self.config.settings.get("output_format", "docx").lower()
        filetypes = "DOCX Files (*.docx);;TXT Files (*.txt)"
        filter_name = "DOCX Files (*.docx)" if output_format == "docx" else "TXT Files (*.txt)"

        save_path, _ = QFileDialog.getSaveFileName(
            parent=None,
            caption="Save Output",
            dir=str(self.config.output_dir / default_name),
            filter=filetypes,
            selectedFilter=filter_name
        )

        if not save_path:
            if status_callback:
                status_callback("Save cancelled by user.", "#ff7373")
            return

        try:
            if save_path.lower().endswith(".txt"):
                combined_text = "\n\n".join([read_file_with_fallback(f) for f in processed_files])
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(combined_text)
            else:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = Document()
                if self.config.settings.get("include_docx_title", True):
                    custom_title = self.config.settings.get("custom_title", "").strip()
                    if not custom_title and video_title:
                        title_text = video_title
                    else:
                        title_text = custom_title or Path(save_path).stem

                    title_para = doc.add_paragraph()
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = title_para.add_run(title_text)
                    run.font.size = Pt(int(self.config.settings.get("title_font_size", 16)))
                    doc.add_paragraph()

                for file in processed_files:
                    content = read_file_with_fallback(file)
                    doc.add_paragraph(content)
                    doc.add_paragraph()
                doc.save(save_path)

            if status_callback:
                status_callback(f"Success: File saved at {save_path}", "#b5e0a8")
        except Exception as e:
            if status_callback:
                status_callback(f"Error saving file: {e}", "#ff7373")
